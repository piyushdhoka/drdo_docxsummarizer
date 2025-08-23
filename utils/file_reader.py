import os
import mimetypes
from typing import Optional, Dict, Any
from pathlib import Path

# Import format-specific readers
from .pdf_reader import extract_text_from_pdf, get_pdf_info

def get_file_info(file_path: str) -> Optional[Dict[str, Any]]:
    """
    Get comprehensive information about any supported file.
    
    Args:
        file_path (str): Path to the file
        
    Returns:
        dict: File information including type, size, pages, etc.
    """
    if not os.path.exists(file_path):
        return None
    
    file_info = {
        'file_name': os.path.basename(file_path),
        'file_size_mb': round(os.path.getsize(file_path) / (1024 * 1024), 2),
        'file_type': get_file_type(file_path),
        'supported': is_file_supported(file_path)
    }
    
    # Add format-specific information
    if file_info['file_type'] == 'pdf':
        pdf_info = get_pdf_info(file_path)
        if pdf_info:
            file_info.update(pdf_info)
    elif file_info['file_type'] == 'docx':
        try:
            import docx
            doc = docx.Document(file_path)
            file_info['page_count'] = len(doc.paragraphs) // 20  # Rough estimate
            file_info['extractor'] = 'python-docx'
        except ImportError:
            file_info['extractor'] = 'python-docx (not installed)'
    
    return file_info

def get_file_type(file_path: str) -> str:
    """Determine file type based on extension and content."""
    ext = Path(file_path).suffix.lower()
    
    if ext == '.pdf':
        return 'pdf'
    elif ext == '.docx':
        return 'docx'
    elif ext == '.txt':
        return 'txt'
    elif ext in ['.html', '.htm']:
        return 'html'
    elif ext == '.md':
        return 'markdown'
    else:
        return 'unknown'

def is_file_supported(file_path: str) -> bool:
    """Check if file format is supported."""
    file_type = get_file_type(file_path)
    return file_type in ['pdf', 'docx', 'txt', 'html', 'markdown']

def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from any supported file format.
    
    Args:
        file_path (str): Path to the file
        
    Returns:
        str: Extracted text or error message
    """
    if not os.path.exists(file_path):
        return f"Error: File '{file_path}' not found"
    
    if not is_file_supported(file_path):
        return f"Error: File format '{get_file_type(file_path)}' is not supported"
    
    file_type = get_file_type(file_path)
    
    try:
        if file_type == 'pdf':
            return extract_text_from_pdf(file_path)
        elif file_type == 'docx':
            return extract_text_from_docx(file_path)
        elif file_type == 'txt':
            return extract_text_from_txt(file_path)
        elif file_type == 'html':
            return extract_text_from_html(file_path)
        elif file_type == 'markdown':
            return extract_text_from_markdown(file_path)
        else:
            return f"Error: Unsupported file type: {file_type}"
    except Exception as e:
        return f"Error extracting text from {file_type.upper()}: {str(e)}"

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX files."""
    try:
        import docx
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text.strip() + "\n\n"
        return text.strip() if text else "Warning: No text could be extracted from the DOCX file."
    except ImportError:
        return "Error: python-docx is not installed. Install it with: pip install python-docx"
    except Exception as e:
        return f"Error: Failed to extract text from DOCX: {str(e)}"

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT files."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            return f"Error: Failed to read TXT file: {str(e)}"
    except Exception as e:
        return f"Error: Failed to read TXT file: {str(e)}"

def extract_text_from_html(file_path: str) -> str:
    """Extract text from HTML files."""
    try:
        from bs4 import BeautifulSoup
        with open(file_path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file.read(), 'html.parser')
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            # Get text and clean it up
            text = soup.get_text()
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            return text
    except ImportError:
        return "Error: beautifulsoup4 is not installed. Install it with: pip install beautifulsoup4"
    except Exception as e:
        return f"Error: Failed to extract text from HTML: {str(e)}"

def extract_text_from_markdown(file_path: str) -> str:
    """Extract text from Markdown files."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            # Simple markdown to text conversion
            import re
            # Remove markdown syntax
            text = re.sub(r'#+\s+', '', content)  # Remove headers
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Remove italic
            text = re.sub(r'`(.*?)`', r'\1', text)  # Remove code
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # Remove links
            return text.strip()
    except Exception as e:
        return f"Error: Failed to read Markdown file: {str(e)}"

def get_supported_formats() -> Dict[str, str]:
    """Get list of supported file formats with descriptions."""
    return {
        'pdf': 'Portable Document Format',
        'docx': 'Microsoft Word Document',
        'txt': 'Plain Text File',
        'html': 'HTML Web Page',
        'markdown': 'Markdown Document'
    }
